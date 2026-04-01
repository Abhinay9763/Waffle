from pathlib import Path
from uuid import uuid4

from fastapi import APIRouter, Depends, File, HTTPException, UploadFile
from fastapi.responses import StreamingResponse
from starlette.status import HTTP_201_CREATED, HTTP_404_NOT_FOUND, HTTP_409_CONFLICT
import httpx
from io import BytesIO
from openpyxl import load_workbook

from deps import get_current_user
from models import QuestionPaper
from supa import db, db_url
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

router = APIRouter()
TEMPLATE_DIR = Path(__file__).resolve().parent.parent / "template"
PAPER_DOCX_TEMPLATE = TEMPLATE_DIR / "Template.docx"


def _normalize_option(value) -> int | None:
    if value is None:
        return None
    s = str(value).strip()
    if not s:
        return None

    # Accept A-D, 0-3, 1-4
    upper = s.upper()
    if upper in {"A", "B", "C", "D"}:
        return {"A": 0, "B": 1, "C": 2, "D": 3}[upper]
    if s in {"0", "1", "2", "3"}:
        return int(s)
    if s in {"1", "2", "3", "4"}:
        return int(s) - 1
    return None


def _parse_import_workbook(raw: bytes):
    try:
        wb = load_workbook(filename=BytesIO(raw), data_only=True)
    except Exception:
        raise HTTPException(
            status_code=400,
            detail={
                "message": "Could not read Excel file.",
                "errors": ["Upload a valid .xlsx template file."],
            },
        )
    ws = wb.active

    rows = list(ws.iter_rows(values_only=True))
    if len(rows) < 2:
        raise HTTPException(status_code=400, detail="Import file has no question rows.")

    headers = [str(h).strip().lower() if h is not None else "" for h in rows[0]]
    header_map = {h: i for i, h in enumerate(headers) if h}

    required = [
        "section", "question", "option_a", "option_b", "option_c", "option_d",
        "correct_option", "marks", "negative_marks",
    ]
    missing = [k for k in required if k not in header_map]
    if missing:
        raise HTTPException(
            status_code=400,
            detail={
                "message": "Missing required columns.",
                "errors": [f"Missing required columns: {', '.join(missing)}"],
            },
        )

    exam_name = "Imported Paper"
    if "exam_name" in header_map:
        for r in rows[1:]:
            candidate = r[header_map["exam_name"]] if header_map["exam_name"] < len(r) else None
            if candidate and str(candidate).strip():
                exam_name = str(candidate).strip()
                break

    sections_by_name: dict[str, dict] = {}
    section_order: list[str] = []
    answers: dict[int, int] = {}
    q_id = 1
    validation_errors: list[str] = []

    last_section_name = ""
    for idx, r in enumerate(rows[1:], start=2):
        section_name = str(r[header_map["section"]]).strip() if header_map["section"] < len(r) and r[header_map["section"]] is not None else ""
        question_text = str(r[header_map["question"]]).strip() if header_map["question"] < len(r) and r[header_map["question"]] is not None else ""

        if not section_name and not question_text:
            continue

        # Allow faculty to leave repeated section cells blank; carry forward previous value.
        if not section_name:
            section_name = last_section_name
        else:
            last_section_name = section_name

        if not section_name:
            validation_errors.append(f"Row {idx}: section is required for the first row of a section block.")
            continue
        if not question_text:
            validation_errors.append(f"Row {idx}: question is required.")
            continue

        options = []
        row_has_option_error = False
        for key in ("option_a", "option_b", "option_c", "option_d"):
            value = r[header_map[key]] if header_map[key] < len(r) else None
            text = "" if value is None else str(value).strip()
            if not text:
                validation_errors.append(f"Row {idx}: {key} is required.")
                row_has_option_error = True
            options.append(text)
        if row_has_option_error:
            continue

        correct_raw = r[header_map["correct_option"]] if header_map["correct_option"] < len(r) else None
        correct_idx = _normalize_option(correct_raw)
        if correct_idx is None:
            validation_errors.append(f"Row {idx}: correct_option must be A-D or 1-4.")
            continue

        marks_raw = r[header_map["marks"]] if header_map["marks"] < len(r) else None
        neg_raw = r[header_map["negative_marks"]] if header_map["negative_marks"] < len(r) else None
        try:
            marks = int(marks_raw)
            negative_marks = int(neg_raw)
        except Exception:
            validation_errors.append(f"Row {idx}: marks/negative_marks must be integers.")
            continue

        if section_name not in sections_by_name:
            section_order.append(section_name)
            sections_by_name[section_name] = {
                "section_id": len(section_order),
                "name": section_name,
                "questions": [],
            }

        sections_by_name[section_name]["questions"].append({
            "question_id": q_id,
            "text": question_text,
            "options": options,
            "correct_option": correct_idx,
            "marks": marks,
            "negative_marks": negative_marks,
        })
        answers[q_id] = correct_idx
        q_id += 1

    sections = [sections_by_name[name] for name in section_order]
    if validation_errors:
        raise HTTPException(
            status_code=400,
            detail={
                "message": "Import validation failed.",
                "errors": validation_errors[:25],
                "error_count": len(validation_errors),
            },
        )

    if not sections:
        raise HTTPException(
            status_code=400,
            detail={
                "message": "No valid question rows were found.",
                "errors": ["Fill at least one valid question row in the template."],
            },
        )

    total_marks = sum(
        q.get("marks", 0)
        for s in sections
        for q in s.get("questions", [])
    )

    return {
        "questions": {
            "meta": {
                "exam_name": exam_name,
                "student_roll": None,
                "start_time": "",
                "end_time": "",
                "total_marks": total_marks,
            },
            "sections": sections,
        },
        "answers": answers,
    }


def collect_image_urls(questions_dict: dict) -> set[str]:
    """Extract every image_url from a questions JSON blob."""
    urls: set[str] = set()
    for section in questions_dict.get("sections", []):
        for q in section.get("questions", []):
            if q.get("image_url"):
                urls.add(q["image_url"])
            for opt in q.get("options", []):
                if isinstance(opt, dict) and opt.get("image_url"):
                    urls.add(opt["image_url"])
    return urls


@router.post("/paper/create", status_code=HTTP_201_CREATED)
async def createPaper(paper: QuestionPaper, user=Depends(get_current_user)):
    data = paper.model_dump(exclude={"id"})
    data["creator_id"] = user["id"]
    response = await db.client.table("QuestionPapers").insert(data).execute()
    return {"msg": "Paper created", "id": response.data[0]["id"]}


@router.post("/paper/import", status_code=HTTP_201_CREATED)
async def importPaper(file: UploadFile = File(...), user=Depends(get_current_user)):
    if not file.filename or not file.filename.lower().endswith(".xlsx"):
        raise HTTPException(status_code=400, detail="Please upload a valid .xlsx file.")

    raw = await file.read()
    if not raw:
        raise HTTPException(status_code=400, detail="Uploaded file is empty.")

    parsed = _parse_import_workbook(raw)
    response = await db.client.table("QuestionPapers").insert({
        "questions": parsed["questions"],
        "answers": parsed["answers"],
        "creator_id": user["id"],
    }).execute()
    return {
        "msg": "Paper imported",
        "id": response.data[0]["id"],
        "question_count": len(parsed["answers"]),
    }


@router.get("/paper/template/{template_name}")
async def downloadTemplate(template_name: str, user=Depends(get_current_user)):
    base = TEMPLATE_DIR.resolve()
    target = (base / template_name).resolve()
    if not str(target).startswith(str(base)):
        raise HTTPException(status_code=400, detail="Invalid template path.")
    if not target.exists() or not target.is_file():
        raise HTTPException(status_code=HTTP_404_NOT_FOUND, detail="Template not found.")

    media = "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
    with target.open("rb") as f:
        data = f.read()
    return StreamingResponse(
        iter([data]),
        media_type=media,
        headers={"Content-Disposition": f"attachment; filename={target.name}"},
    )


@router.get("/paper/list")
async def listPapers(user=Depends(get_current_user)):
    response = await db.client.table("QuestionPapers") \
        .select("id,questions") \
        .eq("creator_id", user["id"]) \
        .execute()

    paper_ids = [p["id"] for p in response.data]
    in_use_ids: set[int] = set()
    if paper_ids:
        exam_res = await db.client.table("Exams") \
            .select("questionpaper_id") \
            .in_("questionpaper_id", paper_ids) \
            .execute()
        in_use_ids = {row["questionpaper_id"] for row in exam_res.data}

    papers = []
    for p in response.data:
        q = p["questions"]
        meta = q.get("meta", {})
        total_marks = meta.get("total_marks") or sum(
            question.get("marks", 0)
            for section in q.get("sections", [])
            for question in section.get("questions", [])
        )
        papers.append({
            "id": p["id"],
            "name": meta.get("exam_name") or f"Paper #{p['id']}",
            "total_marks": total_marks,
            "in_use": p["id"] in in_use_ids,
        })

    return {"papers": papers}


@router.get("/paper/{paper_id:int}")
async def getPaper(paper_id: int, user=Depends(get_current_user)):
    paper_res = await db.client.table("QuestionPapers") \
        .select("*") \
        .eq("id", paper_id) \
        .eq("creator_id", user["id"]) \
        .execute()

    if not paper_res.data:
        raise HTTPException(status_code=HTTP_404_NOT_FOUND, detail="Paper not found.")

    paper = paper_res.data[0]
    exam_res = await db.client.table("Exams") \
        .select("id") \
        .eq("questionpaper_id", paper_id) \
        .limit(1) \
        .execute()
    paper["in_use"] = len(exam_res.data) > 0
    return paper


@router.put("/paper/{paper_id:int}")
async def updatePaper(paper_id: int, paper: QuestionPaper, user=Depends(get_current_user)):
    existing = await db.client.table("QuestionPapers") \
        .select("id,questions") \
        .eq("id", paper_id) \
        .eq("creator_id", user["id"]) \
        .execute()
    if not existing.data:
        raise HTTPException(status_code=HTTP_404_NOT_FOUND, detail="Paper not found.")

    exam_res = await db.client.table("Exams") \
        .select("id") \
        .eq("questionpaper_id", paper_id) \
        .limit(1) \
        .execute()
    if exam_res.data:
        raise HTTPException(status_code=HTTP_409_CONFLICT, detail="Paper is in use by an exam.")

    # Delete images that were removed in this update
    old_urls = collect_image_urls(existing.data[0].get("questions") or {})
    new_urls = collect_image_urls(paper.questions)
    orphaned = old_urls - new_urls
    if orphaned:
        keys = [url.split("/question-images/", 1)[-1] for url in orphaned]
        try:
            await db.client.storage.from_("question-images").remove(keys)
        except Exception:
            pass  # tf should i do

    data = paper.model_dump(exclude={"id", "creator_id"})
    await db.client.table("QuestionPapers").update(data).eq("id", paper_id).execute()
    return {"msg": "Paper updated"}


@router.delete("/paper/{paper_id:int}")
async def deletePaper(paper_id: int, user=Depends(get_current_user)):
    existing = await (db.client.table("QuestionPapers").select("id,questions").eq("id", paper_id).eq("creator_id", user["id"])
                      .execute())
    if not existing.data:
        raise HTTPException(status_code=HTTP_404_NOT_FOUND, detail="Paper not found.")

    exam_res = await db.client.table("Exams").select("id").eq("questionpaper_id", paper_id).limit(1).execute()
    if exam_res.data:
        raise HTTPException(status_code=HTTP_409_CONFLICT, detail="Paper is in use by an exam and cannot be deleted.")

    # Delete all images belonging to this paper
    image_urls = collect_image_urls(existing.data[0].get("questions") or {})
    if image_urls:
        keys = [url.split("/question-images/", 1)[-1] for url in image_urls]
        try:
            await db.client.storage.from_("question-images").remove(keys)
        except Exception:
            pass  # just catch and pass. wtf am i supposed to do here

    await db.client.table("QuestionPapers").delete().eq("id", paper_id).execute()
    return {"msg": "Paper deleted"}

#
# @router.post("/paper/{paper_id}",status_code=HTTP_201_CREATED)
# async def importPaper():
#     pass

@router.post("/paper/{paper_id:int}/clone", status_code=HTTP_201_CREATED)
async def clonePaper(paper_id: int, user=Depends(get_current_user)):
    original = await (db.client.table("QuestionPapers").select("questions,answers").eq("id", paper_id).eq("creator_id", user["id"])
                      .execute())
    if not original.data:
        raise HTTPException(status_code=HTTP_404_NOT_FOUND, detail="Paper not found.")

    src = original.data[0]
    questions = dict(src["questions"])
    meta = dict(questions.get("meta", {}))
    meta["exam_name"] = meta.get("exam_name", "Paper") + " (Copy)"
    questions["meta"] = meta

    result = await db.client.table("QuestionPapers").insert({
        "questions": questions,
        "answers": src["answers"],
        "creator_id": user["id"],
    }).execute()
    return {"msg": "Paper cloned", "id": result.data[0]["id"]}


ALLOWED_IMAGE_TYPES = {"png", "jpg", "jpeg", "gif", "webp"}
QUESTION_IMAGE_WIDTH = 1  # inches
OPTION_IMAGE_WIDTH = 1   # inches

@router.post("/paper/upload-image")
async def uploadImage(file: UploadFile = File(...), user=Depends(get_current_user)):
    ext = (file.filename or "").rsplit(".", 1)[-1].lower()
    if ext not in ALLOWED_IMAGE_TYPES:
        raise HTTPException(status_code=400, detail="Invalid image type. Allowed: png, jpg, jpeg, gif, webp.")
    key = f"{uuid4()}.{ext}"
    data = await file.read()
    await db.client.storage.from_("question-images").upload(key, data, {"content-type": file.content_type})
    url = f"{db_url}/storage/v1/object/public/question-images/{key}"
    return {"url": url}


@router.get("/paper/{paper_id:int}/download")
async def downloadPaperDoc(paper_id: int, user=Depends(get_current_user)):
    """Generate and download a Word document for the question paper with embedded images."""
    # Get the paper
    paper_res = await db.client.table("QuestionPapers") \
        .select("*") \
        .eq("id", paper_id) \
        .eq("creator_id", user["id"]) \
        .execute()

    if not paper_res.data:
        raise HTTPException(status_code=HTTP_404_NOT_FOUND, detail="Paper not found.")

    paper_data = paper_res.data[0]
    questions_data = paper_data.get("questions") or {}
    sections = questions_data.get("sections", [])
    meta = questions_data.get("meta", {})
    paper_name = meta.get("exam_name") or f"Paper #{paper_id}"

    # Prefer the project template and gracefully fall back to the plain format.
    if PAPER_DOCX_TEMPLATE.exists() and PAPER_DOCX_TEMPLATE.is_file():
        doc = Document(str(PAPER_DOCX_TEMPLATE))
    else:
        doc = Document()
    qnum = 0

    # Insert generated content right after the template heading block when possible.
    insert_anchor = None
    for idx, para in enumerate(doc.paragraphs):
        if "www.smec.ac.in" in (para.text or "").lower():
            if idx + 1 < len(doc.paragraphs):
                insert_anchor = doc.paragraphs[idx + 1]
            break

    def add_doc_paragraph(text: str = "", style: str | None = None):
        if insert_anchor is not None:
            return insert_anchor.insert_paragraph_before(text, style=style)
        return doc.add_paragraph(text, style=style)

    # Preserve template body content and append generated questions after it.
    title_style = "Title" if "Title" in doc.styles else None
    section_style = "Heading 1" if "Heading 1" in doc.styles else None
    question_style = "Heading 2" if "Heading 2" in doc.styles else None

    title = add_doc_paragraph(style=title_style)
    title_run = title.add_run(paper_name)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_before = Pt(0)
    if title_style is None:
        title_run.font.size = Pt(16)
        title_run.font.bold = True
    title_run.underline = False

    spacer = add_doc_paragraph()
    spacer.paragraph_format.space_before = Pt(0)
    spacer.paragraph_format.space_after = Pt(0)

    # Fetch images in parallel for better performance
    async with httpx.AsyncClient(timeout=10) as client:
        image_cache = {}

        for section in sections:
            # Add section title if multiple sections
            if len(sections) > 1:
                section_para = add_doc_paragraph(style=section_style)
                section_run = section_para.add_run(section.get("name", "Section"))
                section_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                section_run.underline = True
                if section_style is None:
                    section_run.font.bold = True
                    section_run.font.size = Pt(13)

            for q in section.get("questions", []):
                qnum += 1

                # Question text
                q_para = add_doc_paragraph(style=question_style)
                q_run = q_para.add_run(f"Q{qnum}. {q.get('text', '')}")
                if question_style is None:
                    q_run.font.bold = True

                # Question image
                if q.get("image_url"):
                    try:
                        if q["image_url"] not in image_cache:
                            resp = await client.get(q["image_url"])
                            if resp.status_code == 200:
                                image_cache[q["image_url"]] = resp.content

                        if q["image_url"] in image_cache:
                            img_bytes = BytesIO(image_cache[q["image_url"]])
                            if insert_anchor is not None:
                                img_para = add_doc_paragraph()
                                img_para.add_run().add_picture(img_bytes, width=Inches(QUESTION_IMAGE_WIDTH))
                            else:
                                doc.add_picture(img_bytes, width=Inches(QUESTION_IMAGE_WIDTH))
                    except Exception:
                        pass  # Skip images that fail to download

                # Options
                letters = ["A", "B", "C", "D"]
                for i, opt in enumerate(q.get("options", [])):
                    # Handle both string and OptionValue formats
                    opt_text = opt if isinstance(opt, str) else opt.get("text", "")
                    opt_image_url = None if isinstance(opt, str) else opt.get("image_url")

                    # Option text
                    opt_para = add_doc_paragraph(f"{letters[i]})  {opt_text}")
                    opt_para.paragraph_format.left_indent = Inches(0.3)

                    # Option image
                    if opt_image_url:
                        try:
                            if opt_image_url not in image_cache:
                                resp = await client.get(opt_image_url)
                                if resp.status_code == 200:
                                    image_cache[opt_image_url] = resp.content

                            if opt_image_url in image_cache:
                                img_bytes = BytesIO(image_cache[opt_image_url])
                                if insert_anchor is not None:
                                    opt_img_para = add_doc_paragraph()
                                    opt_img_para.paragraph_format.left_indent = Inches(0.3)
                                    opt_img_para.add_run().add_picture(img_bytes, width=Inches(OPTION_IMAGE_WIDTH))
                                else:
                                    doc.add_picture(img_bytes, width=Inches(OPTION_IMAGE_WIDTH))
                        except Exception:
                            pass  # Skip images that fail to download

    # Write to BytesIO
    output = BytesIO()
    doc.save(output)
    output.seek(0)

    return StreamingResponse(
        iter([output.getvalue()]),
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f"attachment; filename={paper_name}.docx"}
    )
